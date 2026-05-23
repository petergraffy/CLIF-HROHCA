from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


SOURCE = Path("/Users/saborpete/Desktop/Peter/Postdoc/Manuscripts/HROHCA/manuscript_5232026.docx")
OUTPUT = Path("/Users/saborpete/Desktop/Peter/Postdoc/Manuscripts/HROHCA/manuscript_5232026_intro_methods.docx")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_before(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    prior = anchor.insert_paragraph_before("")
    prior.style = style
    prior.add_run(text)
    return prior


def copy_paragraph_style(target: Paragraph, source: Paragraph) -> None:
    target.style = source.style
    target.paragraph_format.left_indent = deepcopy(source.paragraph_format.left_indent)
    target.paragraph_format.right_indent = deepcopy(source.paragraph_format.right_indent)
    target.paragraph_format.first_line_indent = deepcopy(source.paragraph_format.first_line_indent)
    target.paragraph_format.space_before = deepcopy(source.paragraph_format.space_before)
    target.paragraph_format.space_after = deepcopy(source.paragraph_format.space_after)
    target.paragraph_format.line_spacing = deepcopy(source.paragraph_format.line_spacing)


NEW_SECTION = [
    "INTRODUCTION",
    "Extreme heat is an increasingly important cardiovascular threat. Heat exposure has been associated with higher rates of cardiovascular morbidity and mortality, including arrhythmias, ischemic events, heart failure decompensation, and sudden cardiac death. At the same time, climate change is increasing the frequency, intensity, geographic reach, and duration of heat events, making heat-related cardiovascular illness a predictable and recurring challenge for health systems rather than an exceptional disaster scenario.",
    "Out-of-hospital cardiac arrest (OHCA) is among the most severe potential cardiovascular consequences of environmental heat. Prior studies have evaluated temperature and OHCA incidence, emergency medical services activations, or mortality in population-based settings. These studies are essential for defining epidemiologic risk, but they do not fully describe the subset of patients who survive the prehospital phase, achieve return of spontaneous circulation, and require intensive care. This hospitalized subgroup is particularly important for health-system planning because it represents the portion of heat-associated cardiovascular injury that translates directly into ICU beds, mechanical ventilation, vasoactive medications, renal replacement therapy, laboratory monitoring, and complex post-arrest care.",
    "Critical care preparedness is rarely discussed as a central component of climate adaptation. Heat waves can increase acute care demand across multiple clinical syndromes, including cardiovascular disease, heat illness, kidney injury, respiratory failure, and exacerbations of chronic illness. For ICUs, the relevant question is therefore not only whether heat increases the incidence of OHCA, but whether heat-related OHCA (HROHCA) patients who reach the ICU have distinct clinical trajectories or resource needs compared with other OHCA patients. If a reproducible HROHCA phenotype exists, hospitals could use weather forecasts and heat alerts to anticipate post-arrest ICU capacity, staffing, organ support needs, and early monitoring priorities during periods of extreme heat.",
    "The Critical Illness Data Federation (CLIF) provides a multicenter, federated critical care data infrastructure that allows geographically diverse health systems to execute common analytic code locally and share aggregate results without transferring patient-level data. Using this infrastructure, we linked adult ICU admissions for OHCA from 2018 through 2024 to county-level environmental exposure data. Our objectives were twofold: first, to estimate the association between daily maximum temperature and OHCA ICU admission using site-specific distributed lag nonlinear models referenced to local minimum-risk temperature (MRT); and second, to characterize whether HROHCA admissions differ from non-HROHCA admissions in outcomes, organ support, laboratory trajectories, vital sign trajectories, and other markers of ICU resource use.",
    "METHODS",
    "Study Design and Setting",
    "We conducted a retrospective, multicenter federated cohort study across 10 U.S. CLIF sites contributing OHCA, environmental exposure, outcome, and ICU trajectory summaries for calendar years 2018 through 2024. Each site executed a shared R analytic pipeline locally against its CLIF-formatted data. Sites exported aggregate, deidentified CSV files and site-level quality-control figures for central pooling. Patient-level data were not transferred between institutions. The study was designed to evaluate both the temperature-OHCA exposure-response relationship and the clinical phenotype of OHCA patients admitted to an ICU during heat exposure.",
    "Cohort Definition",
    "The study cohort included adult hospitalizations among patients aged 18 years or older with OHCA present on admission and subsequent ICU admission during the index hospitalization. OHCA was identified using present-on-admission cardiac arrest diagnosis codes, including ICD-10 code prefix I46 and ICD-9 code prefix 4275, applied to either hospital diagnosis or admission diagnosis tables depending on local data availability. ICU admission was identified from ADT location records containing ICU-level care. The cohort was limited to admissions from January 1, 2018 through December 31, 2024. Hospitalizations without an assigned county for exposure linkage were excluded from environmental analyses.",
    "County Assignment and Environmental Exposures",
    "Each hospitalization was assigned to a county for exposure linkage using the patient county when available and project-specified hospital-aware reassignment rules when patient county was unavailable or not suitable for federated reporting. Daily county-level maximum temperature (Tmax, degrees Celsius) was derived from Daymet for 2018 through 2024. Daily county-level maximum relative humidity was derived from gridMET. County-year air pollution measures, including nitrogen dioxide (NO2) and fine particulate matter (PM2.5), were retained for secondary pollution-adjusted and pollution-outcome analyses. For daily time-series modeling, site-level daily exposure summaries were constructed from assigned-county environmental values among adult ICU admissions, and daily OHCA counts were modeled over the study period.",
    "Primary Temperature-OHCA Analysis",
    "The primary exposure-response analysis used site-specific distributed lag nonlinear models (DLNMs) to estimate the cumulative association between daily Tmax and OHCA ICU admission. Models used a quasi-Poisson regression framework with a natural spline cross-basis for Tmax over lags 0 through 5 days. The temperature-response dimension used natural splines, and the lag-response dimension used natural splines. Models adjusted for maximum relative humidity, long-term and seasonal time trends, day of week, and calendar year when estimable within site. The primary reference temperature was the site-specific MRT, defined as the temperature on the site-specific prediction grid associated with the lowest estimated cumulative relative risk after initial model fitting. The hot-temperature contrast corresponded to the site-specific 95th percentile of warm-season Tmax.",
    "Site-specific cumulative log relative risks and standard errors were exported for both hot-temperature contrasts and full temperature-response curves. Central pooling used random-effects meta-analysis with DerSimonian-Laird variance estimation. Pooled DLNM curves were generated by pooling site-specific log relative risks at each temperature grid point. Secondary temperature analyses included median-referenced models, humidity-plus-pollution adjusted models, time-adjustment sensitivity analyses, and stratified models by age group, sex, and race where site sample size permitted.",
    "Heat-Related OHCA Phenotype Definition",
    "For clinical phenotyping, HROHCA was defined as warm-season OHCA occurring on a day when the assigned-county Tmax was at or above the site-specific 95th percentile of warm-season Tmax. Non-HROHCA included warm-season OHCA admissions below this site-specific 95th percentile threshold. A less extreme heat90 phenotype, defined using the site-specific 90th percentile of warm-season Tmax, was evaluated as a sensitivity analysis. All phenotype analyses were descriptive and were intended to characterize the ICU burden among OHCA patients who survived to hospital admission rather than to estimate population incidence.",
    "Clinical Outcomes and ICU Trajectories",
    "Clinical outcomes included hospital death, death or hospice discharge, invasive mechanical ventilation (IMV), vasopressor use, ICU length of stay, hospital length of stay, and IMV duration among ventilated patients. ICU resource and trajectory measures included hourly IMV and vasopressor prevalence, CRRT prevalence and initiation windows, cumulative first organ-support events, cumulative discharge and death outcomes, hourly vital signs, hourly laboratory values, and prespecified renal and metabolic marker summaries. Vital signs included heart rate, mean arterial pressure, systolic blood pressure, respiratory rate, oxygen saturation, and temperature. Laboratory trajectories included available chemistry, renal, metabolic, hematologic, and acid-base markers exported by participating sites.",
    "Care Pathways and ICU Timing",
    "To describe the route by which patients reached the ICU, each site summarized admission-to-ICU timing, first hospital location, and observed location pathways derived from ADT records. These summaries were pooled across sites as aggregate counts and percentages. The care pathway analyses were used to contextualize the clinical phenotype and to distinguish patients admitted directly to an ICU from those initially managed in another hospital location before ICU transfer.",
    "Statistical Analysis",
    "Because the analysis was federated, central analyses used aggregate site exports rather than patient-level records. Binary characteristics and outcomes were pooled as summed counts and percentages across sites, with p values based on chi-square or Fisher exact tests when appropriate. Hourly organ-support and cumulative outcome differences were compared using pooled counts and two-sample tests for equality of proportions at each hour. Continuous variables and trajectories were summarized using site medians and interquartile ranges; pooled medians were calculated as sample-size-weighted site medians. Approximate site-level median differences were pooled using random-effects meta-analysis, with site standard errors estimated from the interquartile range and sample size. These continuous trajectory p values should be interpreted as approximate because patient-level distributions were not centrally pooled.",
    "For coefficient-based outcome models, site-specific log odds ratios or log ratios and standard errors were pooled using random-effects meta-analysis. Heterogeneity was summarized using tau-squared, Cochran Q, and I-squared. All analyses were performed using R. Statistical tests were two-sided, and p values were interpreted descriptively in light of multiple trajectory comparisons and the aggregate-data design.",
]


def main() -> None:
    doc = Document(SOURCE)
    intro_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().upper() == "INTRODUCTION")
    results_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().upper() == "RESULTS")
    anchor = doc.paragraphs[results_idx]
    template = doc.paragraphs[intro_idx]

    for text in NEW_SECTION:
        p = insert_paragraph_before(anchor, text)
        copy_paragraph_style(p, template)
        if text in {
            "INTRODUCTION",
            "METHODS",
            "Study Design and Setting",
            "Cohort Definition",
            "County Assignment and Environmental Exposures",
            "Primary Temperature-OHCA Analysis",
            "Heat-Related OHCA Phenotype Definition",
            "Clinical Outcomes and ICU Trajectories",
            "Care Pathways and ICU Timing",
            "Statistical Analysis",
        }:
            for run in p.runs:
                run.bold = True

    # Delete the original Introduction and Methods block after inserting replacement text.
    for p in list(doc.paragraphs[intro_idx:results_idx]):
        delete_paragraph(p)

    consistency_replacements = {
        "across 8 CLIF hospitals": "across 10 CLIF hospitals",
        "Across 8 sites": "Across 10 sites",
        "Across 8 CLIF hospitals": "Across 10 CLIF hospitals",
        "Eight sites contributed DLNM, descriptive, clinical trajectory, and model outputs: Emory, JHU, NU, OHSU, Penn, RUMC, UCMC, and UMN. DLNM estimates were available from all 8 sites.": (
            "Ten sites contributed DLNM, descriptive, clinical trajectory, and model outputs: Emory, JHU, Michigan, NU, OHSU, Penn, RUMC, UCMC, UCSF, and UMN. "
            "DLNM estimates were available from all 10 sites."
        ),
    }
    for paragraph in doc.paragraphs:
      for old, new in consistency_replacements.items():
        if old in paragraph.text:
          paragraph.text = paragraph.text.replace(old, new)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
